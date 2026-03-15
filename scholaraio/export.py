"""
export.py — 论文导出（BibTeX / RIS / Markdown / DOCX 等格式）
==============================================================

将 meta.json 转换为标准引用格式输出。
"""

from __future__ import annotations

import re
from pathlib import Path


def _meta_year(meta: dict) -> int | None:
    """Normalize meta['year'] to int or None, tolerating string values."""
    raw = meta.get("year")
    if raw is None:
        return None
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def _bibtex_escape(text: str) -> str:
    """Escape special LaTeX characters in text."""
    for ch in ("&", "%", "#", "_"):
        text = text.replace(ch, f"\\{ch}")
    return text


def _make_cite_key(meta: dict) -> str:
    """Generate a BibTeX citation key: LastName2023Title."""
    last = meta.get("first_author_lastname") or "Unknown"
    last = re.sub(r"[^a-zA-Z]", "", last)
    year = str(meta.get("year") or "")
    title = meta.get("title") or ""
    # first meaningful word of title (skip short words)
    word = ""
    for w in title.split():
        cleaned = re.sub(r"[^a-zA-Z]", "", w)
        if len(cleaned) > 3:
            word = cleaned.capitalize()
            break
    return f"{last}{year}{word}"


def _type_to_bibtex(paper_type: str) -> str:
    """Map paper_type to BibTeX entry type."""
    mapping = {
        "journal-article": "article",
        "review": "article",
        "book-chapter": "inbook",
        "book": "book",
        "proceedings-article": "inproceedings",
        "conference-paper": "inproceedings",
        "thesis": "phdthesis",
        "dissertation": "phdthesis",
        "preprint": "misc",
        "document": "misc",
        "technical-report": "techreport",
        "manual": "manual",
        "lecture-notes": "misc",
        "standard": "misc",
        "white-paper": "misc",
        "presentation": "misc",
        "meeting-notes": "misc",
    }
    return mapping.get(paper_type or "", "article")


def meta_to_bibtex(meta: dict) -> str:
    """Convert a single meta.json dict to a BibTeX entry string.

    Args:
        meta: Paper metadata dictionary.

    Returns:
        Formatted BibTeX entry string.
    """
    entry_type = _type_to_bibtex(meta.get("paper_type") or "")
    key = _make_cite_key(meta)

    fields: list[tuple[str, str]] = []

    if meta.get("title"):
        fields.append(("title", "{" + _bibtex_escape(meta["title"]) + "}"))
    if meta.get("authors"):
        fields.append(("author", _bibtex_escape(" and ".join(meta["authors"]))))
    if meta.get("year"):
        fields.append(("year", str(meta["year"])))
    if meta.get("journal"):
        fields.append(("journal", _bibtex_escape(meta["journal"])))
    if meta.get("volume"):
        fields.append(("volume", meta["volume"]))
    if meta.get("issue"):
        fields.append(("number", meta["issue"]))
    if meta.get("pages"):
        fields.append(("pages", meta["pages"]))
    if meta.get("publisher"):
        fields.append(("publisher", _bibtex_escape(meta["publisher"])))
    if meta.get("issn"):
        fields.append(("issn", meta["issn"]))
    if meta.get("doi"):
        fields.append(("doi", meta["doi"]))
    if meta.get("abstract"):
        fields.append(("abstract", "{" + _bibtex_escape(meta["abstract"]) + "}"))

    lines = [f"@{entry_type}{{{key},"]
    for name, val in fields:
        lines.append(f"  {name} = {{{val}}},")
    lines.append("}")
    return "\n".join(lines)


def export_bibtex(
    papers_dir: Path,
    *,
    paper_ids: list[str] | None = None,
    year: str | None = None,
    journal: str | None = None,
) -> str:
    """Export papers to BibTeX format.

    Args:
        papers_dir: Root papers directory.
        paper_ids: Specific paper dir names to export. None = all.
        year: Year filter (e.g. "2023", "2020-2024").
        journal: Journal name filter (case-insensitive substring).

    Returns:
        Complete BibTeX string with all matching entries.
    """
    from scholaraio.papers import iter_paper_dirs, parse_year_range, read_meta

    year_start, year_end = parse_year_range(year) if year else (None, None)

    entries: list[str] = []
    for d in iter_paper_dirs(papers_dir):
        if paper_ids and d.name not in paper_ids:
            continue

        meta = read_meta(d)

        # filters
        meta_yr = _meta_year(meta)
        if year_start is not None and (meta_yr or 0) < year_start:
            continue
        if year_end is not None and (meta_yr or 9999) > year_end:
            continue
        if journal and journal.lower() not in (meta.get("journal") or "").lower():
            continue

        entries.append(meta_to_bibtex(meta))

    return "\n\n".join(entries) + "\n" if entries else ""


# ============================================================================
#  RIS export
# ============================================================================

_RIS_TYPE_MAP: dict[str, str] = {
    "journal-article": "JOUR",
    "review": "JOUR",
    "book-chapter": "CHAP",
    "book": "BOOK",
    "proceedings-article": "CONF",
    "conference-paper": "CONF",
    "thesis": "THES",
    "dissertation": "THES",
    "preprint": "UNPB",
    "document": "GEN",
    "technical-report": "RPRT",
    "manual": "MANSCPT",
    "lecture-notes": "GEN",
    "standard": "STAND",
    "white-paper": "RPRT",
    "presentation": "SLIDE",
    "meeting-notes": "GEN",
}


def meta_to_ris(meta: dict) -> str:
    """Convert a single meta.json dict to a RIS entry string.

    Args:
        meta: Paper metadata dictionary.

    Returns:
        Formatted RIS entry string (ends with ``ER  -``).
    """
    ris_type = _RIS_TYPE_MAP.get(meta.get("paper_type") or "", "GEN")
    lines: list[str] = [f"TY  - {ris_type}"]

    if meta.get("title"):
        lines.append(f"TI  - {meta['title']}")
    for author in meta.get("authors") or []:
        lines.append(f"AU  - {author}")
    if meta.get("year"):
        lines.append(f"PY  - {meta['year']}")
    if meta.get("journal"):
        lines.append(f"JO  - {meta['journal']}")
    if meta.get("volume"):
        lines.append(f"VL  - {meta['volume']}")
    if meta.get("issue"):
        lines.append(f"IS  - {meta['issue']}")
    if meta.get("pages"):
        pages = meta["pages"]
        if "-" in str(pages):
            sp, ep = str(pages).split("-", 1)
            lines.append(f"SP  - {sp.strip()}")
            lines.append(f"EP  - {ep.strip()}")
        else:
            lines.append(f"SP  - {pages}")
    if meta.get("publisher"):
        lines.append(f"PB  - {meta['publisher']}")
    if meta.get("issn"):
        lines.append(f"SN  - {meta['issn']}")
    if meta.get("doi"):
        lines.append(f"DO  - {meta['doi']}")
        lines.append(f"UR  - https://doi.org/{meta['doi']}")
    if meta.get("abstract"):
        lines.append(f"AB  - {meta['abstract']}")
    lines.append("ER  -")
    return "\n".join(lines)


def export_ris(
    papers_dir: Path,
    *,
    paper_ids: list[str] | None = None,
    year: str | None = None,
    journal: str | None = None,
) -> str:
    """Export papers to RIS format.

    Args:
        papers_dir: Root papers directory.
        paper_ids: Specific paper dir names to export. None = all.
        year: Year filter (e.g. "2023", "2020-2024").
        journal: Journal name filter (case-insensitive substring).

    Returns:
        Complete RIS string with all matching entries.
    """
    from scholaraio.papers import iter_paper_dirs, parse_year_range, read_meta

    year_start, year_end = parse_year_range(year) if year else (None, None)

    entries: list[str] = []
    for d in iter_paper_dirs(papers_dir):
        if paper_ids and d.name not in paper_ids:
            continue
        meta = read_meta(d)
        meta_yr = _meta_year(meta)
        if year_start is not None and (meta_yr or 0) < year_start:
            continue
        if year_end is not None and (meta_yr or 9999) > year_end:
            continue
        if journal and journal.lower() not in (meta.get("journal") or "").lower():
            continue
        entries.append(meta_to_ris(meta))

    return "\n\n".join(entries) + "\n" if entries else ""


# ============================================================================
#  Markdown reference list export
# ============================================================================


def export_markdown_refs(
    papers_dir: Path,
    *,
    cfg=None,
    paper_ids: list[str] | None = None,
    year: str | None = None,
    journal: str | None = None,
    numbered: bool = True,
    style: str = "apa",
) -> str:
    """Export papers as a Markdown reference list.

    Args:
        papers_dir: Root papers directory.
        cfg: Config object (required for custom styles; built-in styles work without it).
        paper_ids: Specific paper dir names to export. None = all.
        year: Year filter (e.g. "2023", "2020-2024").
        journal: Journal name filter (case-insensitive substring).
        numbered: Use numbered list (default True); False for bullet list.
        style: Citation style name. Built-in: "apa", "vancouver",
            "chicago-author-date", "mla". Custom styles are loaded from
            data/citation_styles/<name>.py. Default: "apa".

    Returns:
        Markdown string with all matching references.
    """
    from scholaraio.citation_styles import BUILTIN_STYLES, FormatterFn, get_formatter
    from scholaraio.papers import iter_paper_dirs, parse_year_range, read_meta

    fmt_fn: FormatterFn
    if style in BUILTIN_STYLES:
        fmt_fn = BUILTIN_STYLES[style]
    else:
        if cfg is None:
            raise ValueError("自定义引用格式需要传入 cfg 参数")
        fmt_fn = get_formatter(style, cfg)

    year_start, year_end = parse_year_range(year) if year else (None, None)

    metas: list[dict] = []
    for d in iter_paper_dirs(papers_dir):
        if paper_ids and d.name not in paper_ids:
            continue
        meta = read_meta(d)
        meta_yr = _meta_year(meta)
        if year_start is not None and (meta_yr or 0) < year_start:
            continue
        if year_end is not None and (meta_yr or 9999) > year_end:
            continue
        if journal and journal.lower() not in (meta.get("journal") or "").lower():
            continue
        metas.append(meta)

    # Sort by year desc, then title
    metas.sort(key=lambda m: (-(_meta_year(m) or 0), m.get("title") or ""))

    lines: list[str] = []
    for i, meta in enumerate(metas, 1):
        lines.append(fmt_fn(meta, i if numbered else None))

    return "\n".join(lines) + "\n" if lines else ""


# ============================================================================
#  DOCX export (general-purpose Markdown → Word)
# ============================================================================


def export_docx(
    content: str,
    output_path: Path,
    *,
    title: str | None = None,
) -> None:
    """Export arbitrary Markdown content to a Word DOCX file.

    Converts Markdown structure (headings, paragraphs, lists, bold/italic,
    tables) to a properly styled Word document. Useful for sharing literature
    reviews, research notes, or any Claude-generated content with colleagues.

    Args:
        content: Markdown text to convert.
        output_path: Destination ``.docx`` file path.
        title: Optional document title (added as Title style if provided).

    Raises:
        ImportError: If ``python-docx`` is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    doc = Document()

    # Set reasonable default styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    if title:
        doc.add_heading(title, level=0)

    _md_to_docx(doc, content)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _md_to_docx(doc, content: str) -> None:
    """Parse Markdown content and add elements to a python-docx Document."""
    try:
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            level = min(level, 9)
            text = line.lstrip("#").strip()
            doc.add_heading(_strip_inline_md(text), level=level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", line):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Fenced code block
        if line.startswith("```") or line.startswith("~~~"):
            fence = line[:3]
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith(fence):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if code_lines:
                p = doc.add_paragraph("\n".join(code_lines))
                try:
                    p.style = doc.styles["No Spacing"]
                except KeyError:
                    pass
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(10)
            continue

        # Table (pipe-separated)
        if "|" in line and line.strip().startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Strip separator row (---)
            data_rows = [r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r.strip())]
            if data_rows:
                cells_per_row = [[c.strip() for c in r.strip().strip("|").split("|")] for r in data_rows]
                ncols = max(len(row) for row in cells_per_row)
                table = doc.add_table(rows=len(cells_per_row), cols=ncols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(cells_per_row):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < ncols:
                            table.cell(ri, ci).text = _strip_inline_md(cell_text)
            continue

        # Unordered list
        if re.match(r"^[\s]*[-*+] ", line):
            text = re.sub(r"^[\s]*[-*+] ", "", line)
            _add_paragraph_with_inline(doc, text, style="List Bullet")
            i += 1
            continue

        # Ordered list
        if re.match(r"^[\s]*\d+\. ", line):
            text = re.sub(r"^[\s]*\d+\. ", "", line)
            _add_paragraph_with_inline(doc, text, style="List Number")
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            p = _add_paragraph_with_inline(doc, text)
            p.paragraph_format.left_indent = Pt(24)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — collect consecutive non-blank, non-special lines
        para_lines: list[str] = []
        while i < len(lines):
            l = lines[i]
            if (
                not l.strip()
                or l.startswith("#")
                or l.startswith("```")
                or l.startswith("~~~")
                or re.match(r"^[-*_]{3,}\s*$", l)
                or ("|" in l and l.strip().startswith("|"))
                or re.match(r"^[\s]*[-*+] ", l)
                or re.match(r"^[\s]*\d+\. ", l)
                or l.startswith("> ")
            ):
                break
            para_lines.append(l)
            i += 1
        if para_lines:
            para_text = " ".join(para_lines)
            _add_paragraph_with_inline(doc, para_text)
        continue

    return


def _strip_inline_md(text: str) -> str:
    """Remove common inline Markdown markers for plain text contexts."""
    # Bold+italic, bold, italic
    text = re.sub(r"\*{3}(.+?)\*{3}", r"\1", text)
    text = re.sub(r"\*{2}(.+?)\*{2}", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # Inline code
    text = re.sub(r"`(.+?)`", r"\1", text)
    # Links
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _add_paragraph_with_inline(doc, text: str, style: str | None = None):
    """Add a paragraph with inline bold/italic formatting preserved. Returns the paragraph."""
    try:
        from docx.shared import Pt
    except ImportError:
        return doc.add_paragraph(text, style=style)

    p = doc.add_paragraph(style=style)
    # Split on bold (**...**) and italic (*...*) patterns
    pattern = re.compile(r"(\*{2,3}.+?\*{2,3}|\*[^*]+\*|__[^_]+__|_[^_]+_|`[^`]+`|\[.+?\]\(.+?\))")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        if re.match(r"\*{3}.+\*{3}", part):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif re.match(r"\*{2}.+\*{2}", part) or re.match(r"__.+__", part):
            run.text = part[2:-2]
            run.bold = True
        elif re.match(r"\*.+\*", part) or re.match(r"_.+_", part):
            run.text = part[1:-1]
            run.italic = True
        elif re.match(r"`[^`]+`", part):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif re.match(r"\[.+?\]\(.+?\)", part):
            # Link — show text only
            m = re.match(r"\[(.+?)\]\((.+?)\)", part)
            if m:
                run.text = m.group(1)
        else:
            run.text = part

    return p
